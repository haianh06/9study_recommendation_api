"""
Script tạo báo cáo DOCX chi tiết về phiên làm việc ngày 23/06/2026.
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, hex_color):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2B579A")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "E8F0FE")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def main():
    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    # ══════════════════════════════════════════════════════════════
    # TRANG BÌA
    # ══════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO KỸ THUẬT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Hệ thống Gợi ý Trường Đại học & Ngành học\n9Study Recommendation Engine")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run("━" * 50)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    run.font.size = Pt(14)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Phiên làm việc: Ngày 23 tháng 06 năm 2026\n")
    run.font.size = Pt(13)
    run = info.add_run("Nội dung: Triển khai phân luồng \"Phòng thi thử\" & \"Khám phá bản thân\"")
    run.font.size = Pt(13)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # MỤC LỤC
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("MỤC LỤC", level=1)
    toc_items = [
        "1. Tổng quan Yêu cầu",
        "2. Phân tích 2 Luồng Hoạt động (User Flow)",
        "3. Các Quyết định Thiết kế Đã Chốt",
        "4. Bảng Mapping Cung Hoàng Đạo",
        "5. Bảng Mapping Thần Số Học",
        "6. Cấu hình Trọng số (weights.json)",
        "7. Kiến trúc Mã nguồn & Các Thay đổi",
        "8. API Endpoint - Cập nhật",
        "9. Sửa lỗi (Bug Fixes)",
        "10. Kiểm thử & Kết quả",
        "11. Danh sách File Đã Thay đổi"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. TỔNG QUAN YÊU CẦU
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. Tổng quan Yêu cầu", level=1)
    doc.add_paragraph(
        "Hệ thống 9Study Recommendation Engine là một công cụ gợi ý Trường Đại học và Ngành học "
        "dành cho học sinh ôn thi THPT Quốc Gia. Trước phiên làm việc này, hệ thống đã có sẵn "
        "các chức năng cơ bản: xét tuyển theo khối thi, điểm mục tiêu, sở thích cá nhân và nhóm "
        "tính cách Holland/MBTI."
    )
    doc.add_paragraph(
        "Yêu cầu của phiên làm việc ngày hôm nay là triển khai 2 luồng hoạt động mới cho người dùng, "
        "tích hợp thêm yếu tố Thần số học (Numerology) và Cung hoàng đạo (Zodiac) vào thuật toán đánh giá "
        "mức độ phù hợp giữa người dùng và các ngành học."
    )

    # ══════════════════════════════════════════════════════════════
    # 2. PHÂN TÍCH 2 LUỒNG
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. Phân tích 2 Luồng Hoạt động (User Flow)", level=1)

    doc.add_heading("2.1. Điểm chung của cả 2 luồng", level=2)
    common_steps = [
        "Chọn bài thi (Khối thi: A00, A01, B00, D01,...)",
        "Chọn Nguyện vọng thi (Trường, Ngành học) - có thể nhiều trường, nhiều ngành",
        "Nhập Điểm mục tiêu (thang 30)",
        "Nhập Điểm trung bình hiện tại",
    ]
    for step in common_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading("2.2. Luồng 1: Phòng thi thử (Mock Exam)", level=2)
    doc.add_paragraph(
        "Sau khi nhập Điểm trung bình hiện tại, người dùng nhập Ngày sinh (DOB). "
        "Hệ thống sử dụng DOB để tự động tính Thần số học và Cung hoàng đạo, "
        "sau đó kết hợp với dữ liệu sở thích (Data/Interest) để tính tỷ lệ phần trăm "
        "người dùng phù hợp với các ngành học đã chọn."
    )
    doc.add_paragraph("Input đặc thù: current_score (Điểm TB hiện tại) + dob (Ngày sinh).", style='List Bullet')
    doc.add_paragraph("Không có bài test Holland.", style='List Bullet')
    doc.add_paragraph("Điểm TB hiện tại được dùng để tính recommendation_score (mức độ phù hợp về điểm chuẩn).", style='List Bullet')

    doc.add_heading("2.3. Luồng 2: Khám phá bản thân (Discovery)", level=2)
    doc.add_paragraph(
        "Sau khi nhập Điểm mục tiêu, người dùng nhập Ngày sinh. "
        "Người dùng KHÔNG nhập Điểm trung bình hiện tại trong luồng này. "
        "Tiếp theo, người dùng có thể chọn làm hoặc không làm bài test Holland (RIASEC)."
    )

    doc.add_heading("2.3a. Luồng 2a: Không làm test Holland", level=3)
    doc.add_paragraph("Đánh giá dựa trên: Data/Sở thích + Thần số học + Cung hoàng đạo.")

    doc.add_heading("2.3b. Luồng 2b: Có làm test Holland", level=3)
    doc.add_paragraph("Đánh giá dựa trên: Data/Sở thích + Holland (RIASEC) + Thần số học + Cung hoàng đạo.")

    doc.add_heading("2.4. Unhappy Case (Không có Nguyện vọng)", level=2)
    doc.add_paragraph(
        "Nếu người dùng không nhập bất kỳ nguyện vọng nào (aspirations rỗng), hệ thống sẽ tự động "
        "chuyển sang chế độ Cold-Start: trả về danh sách top_majors (các nhóm ngành phù hợp nhất) "
        "kèm phần trăm match, và đề xuất các chương trình đào tạo thuộc các nhóm ngành đó."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. CÁC QUYẾT ĐỊNH THIẾT KẾ ĐÃ CHỐT
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. Các Quyết định Thiết kế Đã Chốt (User Approved)", level=1)

    decisions = [
        ["Điểm TB Hiện tại (current_score)", "Chỉ có ở Luồng 1 (Phòng thi thử). Được sử dụng duy nhất để tính recommendation_score (xác định khả năng đỗ: An toàn / Vừa sức / Thử thách)."],
        ["Tính toán Thần số học & Cung hoàng đạo", "Backend tự động tính từ chuỗi DOB (YYYY-MM-DD). Frontend chỉ cần gửi ngày sinh, không cần tính toán phía client."],
        ["Cơ chế trọng số động", "Trọng số của các yếu tố được tách ra file weights.json để dễ dàng điều chỉnh mà không cần sửa code Python."],
        ["Bảng Mapping Cung Hoàng Đạo", "Sử dụng bảng tiêu chuẩn chia theo 4 nhóm nguyên tố (Lửa, Đất, Khí, Nước), mỗi cung ánh xạ tới 4 nhóm ngành phù hợp."],
        ["Bỏ trường preferred_provinces", "Loại bỏ hoàn toàn trường Tỉnh/TP ưu tiên học ra khỏi API và thuật toán."]
    ]
    add_styled_table(doc, ["Quyết định", "Chi tiết"], decisions, [5, 12])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. BẢNG MAPPING CUNG HOÀNG ĐẠO
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. Bảng Mapping Cung Hoàng Đạo (ZODIAC_MAP)", level=1)
    doc.add_paragraph(
        "Bảng mapping dưới đây ánh xạ 12 Cung hoàng đạo với các nhóm ngành học phù hợp. "
        "Được chia theo 4 nhóm nguyên tố: Lửa (nhiệt huyết), Đất (thực tế), Khí (giao tiếp, sáng tạo), Nước (cảm xúc, chữa lành)."
    )

    zodiac_rows = [
        ["♈ Lửa", "Aries\n(Bạch Dương)\n21/3 - 19/4", "Quân sự - Công an\nKinh tế - QTKD\nThể dục - Thể thao\nKỹ thuật ô tô - Cơ khí"],
        ["♈ Lửa", "Leo\n(Sư Tử)\n23/7 - 22/8", "Kinh tế - QTKD\nNghệ thuật - Thiết kế\nDu lịch - Khách sạn\nLuật - Pháp lý"],
        ["♈ Lửa", "Sagittarius\n(Nhân Mã)\n22/11 - 21/12", "Du lịch - Khách sạn\nNgôn ngữ - Ngoại ngữ\nSư phạm\nKinh doanh quốc tế"],
        ["♉ Đất", "Taurus\n(Kim Ngưu)\n20/4 - 20/5", "Kế toán - Kiểm toán\nTài chính - Ngân hàng\nNông nghiệp - BVTV\nNghệ thuật - Thiết kế"],
        ["♉ Đất", "Virgo\n(Xử Nữ)\n23/8 - 22/9", "KHMT - KTPM\nKế toán - Kiểm toán\nY - Dược\nToán - Thống kê ứng dụng"],
        ["♉ Đất", "Capricorn\n(Ma Kết)\n22/12 - 19/1", "Tài chính - Ngân hàng\nXây dựng - Kiến trúc\nQLNN - Nhân lực\nKế toán - Kiểm toán"],
        ["♊ Khí", "Gemini\n(Song Tử)\n21/5 - 20/6", "Ngôn ngữ - Ngoại ngữ\nMarketing - PR\nCNTT - Truyền thông\nDu lịch - Khách sạn"],
        ["♊ Khí", "Libra\n(Thiên Bình)\n23/9 - 22/10", "Luật - Pháp lý\nNghệ thuật - Thiết kế\nMarketing - PR\nVăn học - KHXH"],
        ["♊ Khí", "Aquarius\n(Bảo Bình)\n20/1 - 18/2", "CNTT - Truyền thông\nKHMT - KTPM\nĐiện - Điện tử - TĐH\nKHTN - Vật lý - Hóa học"],
        ["♋ Nước", "Cancer\n(Cự Giải)\n21/6 - 22/7", "Y - Dược\nTâm lý học - CTXH\nSư phạm\nQLNN - Nhân lực"],
        ["♋ Nước", "Scorpio\n(Thiên Yết)\n23/10 - 21/11", "Tâm lý học - CTXH\nY - Dược\nKHTN - Vật lý - Hóa học\nKinh doanh quốc tế"],
        ["♋ Nước", "Pisces\n(Song Ngư)\n19/2 - 20/3", "Nghệ thuật - Thiết kế\nTâm lý học - CTXH\nY - Dược\nVăn học - KHXH"],
    ]
    add_styled_table(doc, ["Nhóm", "Cung hoàng đạo", "Nhóm ngành phù hợp"], zodiac_rows, [2, 4, 11])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. BẢNG MAPPING THẦN SỐ HỌC
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. Bảng Mapping Thần Số Học (NUMEROLOGY_MAP)", level=1)
    doc.add_paragraph(
        "Bảng dưới đây ánh xạ Số chủ đạo (Life Path Number) được tính toán từ Ngày sinh "
        "với các Nhóm ngành phù hợp. Bao gồm 9 số cơ bản (1-9) và 3 Master Number (11, 22, 33)."
    )

    numerology_rows = [
        ["1", "Người Lãnh đạo", "Kinh tế - QTKD, Kinh doanh quốc tế, Quân sự - Công an, CNTT - Truyền thông"],
        ["2", "Người Hợp tác", "Tâm lý học - CTXH, Sư phạm, Ngôn ngữ - Ngoại ngữ, Y - Dược"],
        ["3", "Người Sáng tạo", "Nghệ thuật - Thiết kế, Marketing - PR, Du lịch - Khách sạn, Văn học - KHXH"],
        ["4", "Người Xây dựng", "Kế toán - Kiểm toán, Tài chính - Ngân hàng, Xây dựng - Kiến trúc, Luật - Pháp lý"],
        ["5", "Người Tự do", "Du lịch - Khách sạn, Marketing - PR, Ngôn ngữ - Ngoại ngữ, Kinh doanh quốc tế"],
        ["6", "Người Chữa lành", "Y - Dược, Sư phạm, Tâm lý học - CTXH, Công nghệ thực phẩm"],
        ["7", "Người Tìm kiếm", "KHMT - KTPM, KHTN - Vật lý - Hóa học, Địa lý - Môi trường, Toán - Thống kê"],
        ["8", "Người Quyền lực", "Tài chính - Ngân hàng, Kinh tế - QTKD, Kế toán - Kiểm toán, Luật - Pháp lý"],
        ["9", "Người Nhân ái", "Sư phạm, Y - Dược, Tâm lý học - CTXH, Nghệ thuật - Thiết kế"],
        ["11", "Master: Trực giác", "Tâm lý học - CTXH, Văn học - KHXH, Nghệ thuật - Thiết kế, Sư phạm"],
        ["22", "Master: Kiến tạo", "Xây dựng - Kiến trúc, KHMT - KTPM, Kinh tế - QTKD, Công nghệ sinh học"],
        ["33", "Master: Thầy giáo", "Sư phạm, Y - Dược, Tâm lý học - CTXH, Nghệ thuật - Thiết kế"],
    ]
    add_styled_table(doc, ["Số", "Đặc trưng", "Nhóm ngành phù hợp"], numerology_rows, [1.5, 4, 11.5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. CẤU HÌNH TRỌNG SỐ
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. Cấu hình Trọng số (weights.json)", level=1)
    doc.add_paragraph(
        "File weights.json chứa trọng số (%) của từng yếu tố tham gia vào công thức tính "
        "match_percentage (tỷ lệ phù hợp giữa người dùng và nhóm ngành). "
        "File này nằm ở thư mục gốc dự án, có thể chỉnh sửa trực tiếp mà không cần sửa code Python."
    )

    weight_rows = [
        ["mock_exam\n(Phòng thi thử)", "40%", "—", "30%", "30%"],
        ["discovery_no_holland\n(Khám phá - Không Holland)", "40%", "—", "30%", "30%"],
        ["discovery_holland\n(Khám phá - Có Holland)", "55%", "15%", "15%", "15%"],
    ]
    add_styled_table(doc, ["Luồng (Flow)", "Data / Sở thích", "Holland", "Thần số học", "Cung hoàng đạo"], weight_rows, [5, 3, 2.5, 3, 3.5])

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Lưu ý: ")
    run.bold = True
    p.add_run("Để thay đổi trọng số, chỉ cần mở file weights.json và chỉnh sửa giá trị. "
              "Hệ thống sẽ tự động đọc lại file này khi khởi động lại server.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. KIẾN TRÚC MÃ NGUỒN & CÁC THAY ĐỔI
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. Kiến trúc Mã nguồn & Các Thay đổi", level=1)

    doc.add_heading("7.1. personality_mapper.py", level=2)
    changes_pm = [
        ["ZODIAC_MAP", "Thêm mới", "Bộ từ điển ánh xạ 12 cung hoàng đạo tới các nhóm ngành tương ứng."],
        ["calculate_numerology(dob_str)", "Thêm mới", "Tính Số chủ đạo (Life Path Number) từ ngày sinh. Hỗ trợ Master Number (11, 22, 33)."],
        ["get_zodiac_sign(dob_str)", "Thêm mới", "Xác định Cung hoàng đạo từ ngày sinh (định dạng YYYY-MM-DD)."],
    ]
    add_styled_table(doc, ["Thành phần", "Loại thay đổi", "Mô tả"], changes_pm, [5, 3, 9])

    doc.add_heading("7.2. recommender.py", level=2)
    changes_rec = [
        ["UserProfile", "Sửa đổi", "Thêm trường flow_type, zodiac, current_score. Loại bỏ preferred_provinces."],
        ["__init__()", "Sửa đổi", "Bổ sung hàm _load_flow_weights() để đọc weights.json khi khởi tạo engine."],
        ["_calculate_major_match_percent()", "Sửa đổi", "Tích hợp zodiac_score từ ZODIAC_MAP. Sử dụng trọng số động từ self.flow_weights thay vì hằng số cứng."],
        ["_compute_scores()", "Sửa đổi", "Sử dụng current_score (nếu có) thay cho total_score khi tính _score_match."],
        ["_location_match()", "Sửa đổi", "Đơn giản hóa, loại bỏ preferred_provinces."],
        ["recommend()", "Sửa đổi", "Phân tách luồng đầu ra: tính % cho từng nguyện vọng → aspiration_matches. Nếu rỗng → top_majors."],
        ["print_recommendations()", "Thêm mới", "Hàm helper in kết quả đẹp ra console cho demo/debug."],
    ]
    add_styled_table(doc, ["Thành phần", "Loại thay đổi", "Mô tả"], changes_rec, [5, 3, 9])

    doc.add_heading("7.3. api.py", level=2)
    changes_api = [
        ["RecommendationRequest", "Sửa đổi", "Thêm flow_type, dob, current_score. Loại bỏ preferred_provinces. Cập nhật Example Payload cho Swagger UI."],
        ["AspirationMatchResponse", "Thêm mới", "Model response cho kết quả % phù hợp của từng nguyện vọng."],
        ["RecommendationResponse", "Sửa đổi", "Thêm trường aspiration_matches."],
        ["get_recommendations()", "Sửa đổi", "Tiền xử lý DOB → numerology + zodiac trước khi build UserProfile."],
    ]
    add_styled_table(doc, ["Thành phần", "Loại thay đổi", "Mô tả"], changes_api, [5, 3, 9])

    doc.add_heading("7.4. Các file mới", level=2)
    new_files = [
        ["weights.json", "File cấu hình trọng số cho từng luồng (mock_exam, discovery_no_holland, discovery_holland)."],
        ["requirements.txt", "Danh sách tất cả các thư viện Python cần thiết để chạy dự án."],
    ]
    add_styled_table(doc, ["File", "Mô tả"], new_files, [5, 12])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. API ENDPOINT - CẬP NHẬT
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. API Endpoint - Cập nhật", level=1)

    doc.add_heading("8.1. POST /recommend - Request Body", level=2)
    doc.add_paragraph("Các trường mới được thêm vào Request Body (ngoài các trường hiện có):")

    api_fields = [
        ["flow_type", "string", '"mock_exam"', "Bắt buộc", '"mock_exam" hoặc "discovery"'],
        ["current_score", "float", "null", "Tuỳ chọn", "Điểm TB hiện tại (0-30). Chỉ dùng ở luồng Phòng thi thử."],
        ["dob", "string", "null", "Tuỳ chọn", 'Ngày sinh định dạng "YYYY-MM-DD". Backend tự tính numerology & zodiac.'],
        ["aspirations", "array", "[]", "Tuỳ chọn", "Danh sách nguyện vọng [{university_code, major_name}]."],
    ]
    add_styled_table(doc, ["Trường", "Kiểu", "Mặc định", "Bắt buộc?", "Mô tả"], api_fields, [3, 2, 2, 2, 8])

    doc.add_heading("8.2. POST /recommend - Response Body", level=2)
    doc.add_paragraph("Các trường mới trong Response:")

    api_resp = [
        ["aspiration_matches", "array", "Danh sách [{university_code, major_name, major_group_name, match_percentage}] cho từng nguyện vọng."],
        ["top_majors", "array", "Chỉ trả về khi aspirations rỗng. Danh sách [{major_group_name, match_percentage}]."],
    ]
    add_styled_table(doc, ["Trường", "Kiểu", "Mô tả"], api_resp, [4, 2, 11])

    doc.add_heading("8.3. Ví dụ Request Payload", level=2)

    example_payloads = [
        ("Luồng 1: Phòng thi thử", '''{
  "flow_type": "mock_exam",
  "exam_block": ["A00", "A01"],
  "total_score": 26.5,
  "current_score": 24.0,
  "province": "Hà Nội",
  "dob": "2006-05-12",
  "aspirations": [
    {"university_code": "BKA", "major_name": "Khoa học Máy tính"},
    {"university_code": "QGH", "major_name": "Công nghệ thông tin"},
    {"university_code": "FPT", "major_name": "Kỹ thuật phần mềm"}
  ]
}'''),
        ("Luồng 2a: Khám phá - Không Holland", '''{
  "flow_type": "discovery",
  "exam_block": ["D01"],
  "total_score": 25.0,
  "province": "TP.HCM",
  "dob": "2006-10-20",
  "aspirations": [
    {"university_code": "UEH", "major_name": "Kinh doanh quốc tế"},
    {"university_code": "FTU", "major_name": "Quản trị kinh doanh"},
    {"university_code": "VNU", "major_name": "Ngôn ngữ Anh"}
  ]
}'''),
        ("Luồng 2b: Khám phá - Có Holland", '''{
  "flow_type": "discovery",
  "exam_block": ["B00"],
  "total_score": 27.5,
  "province": "Đà Nẵng",
  "dob": "2006-02-14",
  "holland_code": "SIA",
  "aspirations": [
    {"university_code": "YDS", "major_name": "Y khoa"},
    {"university_code": "YHN", "major_name": "Răng - Hàm - Mặt"},
    {"university_code": "VNU", "major_name": "Tâm lý học"}
  ]
}'''),
        ("Unhappy Case: Không có Nguyện vọng", '''{
  "flow_type": "mock_exam",
  "exam_block": ["A01"],
  "total_score": 22.0,
  "current_score": 20.5,
  "province": "Hải Phòng",
  "dob": "2006-12-05",
  "aspirations": []
}'''),
    ]

    for title, code in example_payloads:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. SỬA LỖI (BUG FIXES)
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("9. Sửa lỗi (Bug Fixes)", level=1)

    bugs = [
        ["program_id trả về null", "data_loader.py", "Nghiêm trọng",
         "Khi thực hiện merge dữ liệu từ các bảng (admission_scores, program_admissions), "
         "Pandas tạo ra cột program_id trùng lặp. Cột chứa NaN bị sử dụng thay vì cột gốc.",
         "Drop cột program_id thừa ngay sau mỗi bước merge, trước khi rename cột id cuối cùng."],
        ["print_recommendations bị thiếu", "recommender.py", "Trung bình",
         "Hàm print_recommendations() không tồn tại trong recommender.py khiến demo.py crash khi import.",
         "Thêm lại hàm print_recommendations() vào cuối file recommender.py."],
        ["UnicodeEncodeError khi chạy trên Windows", "api.py, db_introspection.py", "Nhẹ",
         "Các ký tự emoji (🔄, ✅, 🔌) không encode được với charset cp1252 mặc định của Windows.",
         "Đã ghi nhận. Có thể khắc phục bằng cách thêm sys.stdout.reconfigure(encoding='utf-8') hoặc thay emoji bằng ký tự ASCII."],
    ]
    add_styled_table(doc, ["Lỗi", "File", "Mức độ", "Nguyên nhân", "Cách khắc phục"], bugs, [3, 2.5, 2, 4.5, 5])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. KIỂM THỬ & KẾT QUẢ
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("10. Kiểm thử & Kết quả", level=1)
    doc.add_paragraph(
        "Toàn bộ 6 kịch bản test (4 kịch bản cũ + 2 kịch bản mới cho luồng Phòng thi thử và Khám phá bản thân) "
        "đã được tích hợp vào file demo.py duy nhất và chạy thành công không lỗi."
    )

    test_results = [
        ["Test 1", "Học sinh khá - A00 - Hà Nội", "4.715", "10", "match: 9, safe: 1"],
        ["Test 2", "Học sinh giỏi - D01 - TP.HCM", "5.272", "10", "match: 8, safe: 2"],
        ["Test 3", "Cold Start - Chỉ thông tin cơ bản", "4.445", "10", "safe: 9, match: 1"],
        ["Test 4", "Ngành Y Dược - Điểm cao", "3.327", "10", "match: 8, safe: 2"],
        ["Test 5", "Flow Phòng thi thử + Nguyện vọng", "4.817", "5", "unknown: 3, safe: 2"],
        ["Test 6", "Flow Khám phá + Holland", "4.709", "5", "match: 4, safe: 1"],
    ]
    add_styled_table(doc, ["#", "Kịch bản", "Candidates", "Kết quả", "Safety Distribution"], test_results, [1.5, 6, 2.5, 2, 5])

    doc.add_paragraph()

    doc.add_heading("10.1. Kết quả Aspiration Matches (Test Case 5 & 6)", level=2)

    aspiration_results = [
        ["Test 5", "Công nghệ thông tin", "CNTT - Truyền thông", "60%"],
        ["Test 5", "Kinh tế quốc tế", "Kinh tế - QTKD", "60%"],
        ["Test 6", "Tâm lý học", "Tâm lý học - CTXH", "100%"],
    ]
    add_styled_table(doc, ["Test", "Nguyện vọng", "Nhóm ngành mapping", "% Phù hợp"], aspiration_results, [2, 4, 5, 3])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 11. DANH SÁCH FILE ĐÃ THAY ĐỔI
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("11. Danh sách File Đã Thay đổi", level=1)

    file_list = [
        ["personality_mapper.py", "MODIFY", "Thêm ZODIAC_MAP, calculate_numerology(), get_zodiac_sign()."],
        ["recommender.py", "MODIFY", "Cập nhật UserProfile, thuật toán tính điểm, trọng số động, print_recommendations()."],
        ["api.py", "MODIFY", "Thêm flow_type, dob, current_score. Loại bỏ preferred_provinces. Cập nhật Swagger Example."],
        ["data_loader.py", "MODIFY", "Fix bug duplicate program_id khi merge."],
        ["demo.py", "MODIFY", "Tích hợp Test Case 5 & 6 cho luồng mới. Loại bỏ preferred_provinces."],
        ["weights.json", "NEW", "File cấu hình trọng số cho từng luồng."],
        ["requirements.txt", "NEW", "Danh sách thư viện Python của dự án."],
        ["test_flows.py", "DELETE", "Đã tích hợp toàn bộ vào demo.py. File đã bị xóa."],
    ]
    add_styled_table(doc, ["File", "Loại", "Ghi chú"], file_list, [4.5, 2, 10.5])

    # ══════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════
    output_path = os.path.join(os.path.dirname(__file__), "9study_Daily_Report_20260623.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    main()
