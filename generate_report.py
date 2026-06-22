"""
Generate comprehensive DOCX report for the 9study project.
"""
import os
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# ──────────────────────────────────────────
# Helper Functions
# ──────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a nicely formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2E4057")
    
    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_shading(cell, "F0F4F8")
    
    return table


def add_code_block(doc, code_text, language="python"):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add background shading to the paragraph
    pPr = p._element.get_or_add_pPr()
    shading = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F5F5F5',
        qn('w:val'): 'clear',
    })
    pPr.append(shading)


def add_section_break(doc):
    """Add visual separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


# ──────────────────────────────────────────
# Read source files
# ──────────────────────────────────────────

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))

def read_file(filename):
    path = os.path.join(PROJECT_DIR, filename)
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return f"[File not found: {filename}]"

source_files = {
    "api.py": read_file("api.py"),
    "recommender.py": read_file("recommender.py"),
    "data_loader.py": read_file("data_loader.py"),
    "personality_mapper.py": read_file("personality_mapper.py"),
    "evaluate.py": read_file("evaluate.py"),
    "demo.py": read_file("demo.py"),
    "db_introspection.py": read_file("db_introspection.py"),
    "create_eda_notebook.py": read_file("create_eda_notebook.py"),
    "requirements.txt": read_file("requirements.txt"),
    ".gitignore": read_file(".gitignore"),
}


# ──────────────────────────────────────────
# Build Document
# ──────────────────────────────────────────

doc = Document()

# Configure styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

# ═══════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BÁO CÁO PHÂN TÍCH MÃ NGUỒN")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Dự án 9study — Hệ thống Gợi ý Ngành học & Trường Đại học")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x5A, 0x6A, 0x7A)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Career & University Recommendation System")
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("Ngày tạo: 21/06/2026")
run.font.size = Pt(12)

version_p = doc.add_paragraph()
version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version_p.add_run("Phiên bản: 1.0.0")
run.font.size = Pt(12)

doc.add_page_break()


# ═══════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════

doc.add_heading("MỤC LỤC", level=1)

toc_items = [
    "1. Tổng quan dự án",
    "2. Kiến trúc hệ thống", 
    "3. Cấu trúc thư mục & Mô tả file",
    "4. Chi tiết các Module",
    "   4.1. data_loader.py — Module Tải & Tiền xử lý Dữ liệu",
    "   4.2. recommender.py — Engine Gợi ý (Core)",
    "   4.3. personality_mapper.py — Ánh xạ Tính cách",
    "   4.4. api.py — FastAPI REST API",
    "   4.5. evaluate.py — Đánh giá Hệ thống",
    "   4.6. demo.py — Script Demo",
    "   4.7. db_introspection.py — Khảo sát Database",
    "   4.8. create_eda_notebook.py — Tạo Notebook EDA",
    "5. Mô hình dữ liệu (Data Model)",
    "6. Thuật toán Gợi ý — Pipeline 3 Giai đoạn",
    "7. API Reference",
    "8. Cấu hình & Phụ thuộc (Dependencies)",
    "9. Đánh giá & Metrics",
    "10. Mã nguồn đầy đủ (Full Source Code)",
]

for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    if not item.startswith("   "):
        run.bold = True

doc.add_page_break()


# ═══════════════════════════════════════════
# 1. TỔNG QUAN DỰ ÁN
# ═══════════════════════════════════════════

doc.add_heading("1. Tổng quan dự án", level=1)

doc.add_paragraph(
    "9study Career & University Recommendation System là một hệ thống gợi ý thông minh "
    "giúp học sinh THPT tìm kiếm ngành học và trường đại học phù hợp nhất dựa trên hồ sơ cá nhân. "
    "Hệ thống sử dụng phương pháp kết hợp giữa Knowledge-Based Filtering và Content-Based Scoring "
    "để đưa ra các đề xuất được cá nhân hóa."
)

doc.add_heading("Mục tiêu chính", level=2)
objectives = [
    "Gợi ý ngành học và trường đại học phù hợp dựa trên điểm thi, khối thi, sở thích và ngân sách",
    "Hỗ trợ Cold-Start thông qua ánh xạ tính cách (MBTI, Holland RIASEC, Thần số học)",
    "Cung cấp REST API để tích hợp với frontend (web/mobile)",
    "Phân loại mức độ an toàn (Safe/Match/Reach/Ambitious) cho từng đề xuất",
    "Đảm bảo đa dạng trong kết quả gợi ý (trường, ngành, khu vực)",
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading("Công nghệ sử dụng", level=2)
tech_data = [
    ["Ngôn ngữ", "Python 3.12+"],
    ["Web Framework", "FastAPI + Uvicorn"],
    ["Database", "PostgreSQL (SQLAlchemy ORM)"],
    ["Xử lý dữ liệu", "Pandas, NumPy"],
    ["Trực quan hóa", "Matplotlib, Seaborn"],
    ["ML/Scoring", "Scikit-learn, Custom Scoring Engine"],
    ["API Documentation", "Swagger UI (tự động từ FastAPI)"],
    ["Environment", "python-dotenv"],
]
add_styled_table(doc, ["Thành phần", "Công nghệ"], tech_data)

doc.add_page_break()


# ═══════════════════════════════════════════
# 2. KIẾN TRÚC HỆ THỐNG
# ═══════════════════════════════════════════

doc.add_heading("2. Kiến trúc hệ thống", level=1)

doc.add_paragraph(
    "Hệ thống được thiết kế theo kiến trúc layered (phân tầng) với luồng dữ liệu rõ ràng "
    "từ Database → Data Loader → Recommendation Engine → REST API → Client."
)

doc.add_heading("Sơ đồ kiến trúc tổng thể", level=2)

arch_text = """
┌──────────────────────────────────────────────────────────────────┐
│                        CLIENT (Web / Mobile)                     │
│                      POST /recommend, GET /health                │
└──────────────────────────┬───────────────────────────────────────┘
                           │ HTTP/JSON
┌──────────────────────────▼───────────────────────────────────────┐
│                      API Layer (api.py)                           │
│                FastAPI + Pydantic Validation + CORS               │
│   Endpoints: /recommend, /health, /major-groups, /personality     │
└──────────────────────────┬───────────────────────────────────────┘
                           │
┌──────────────────────────▼───────────────────────────────────────┐
│                 Recommendation Engine (recommender.py)            │
│                                                                   │
│  ┌─────────────┐   ┌──────────────┐   ┌────────────────────┐    │
│  │  Stage 1:   │──>│   Stage 2:   │──>│    Stage 3:        │    │
│  │ Hard Filter │   │ Content-Based│   │ Re-rank + Safety   │    │
│  │ (Rules)     │   │   Scoring    │   │  Classification    │    │
│  └─────────────┘   └──────────────┘   └────────────────────┘    │
│                                                                   │
│  Inputs: UserProfile (exam_block, score, province, interests...) │
│  Output: Ranked list + Top Majors + Safety Labels                 │
└──────────────────────────┬───────────────────────────────────────┘
                           │
┌──────────────────────────▼───────────────────────────────────────┐
│                  Data Layer                                       │
│  ┌──────────────────┐  ┌────────────────────────────┐            │
│  │  data_loader.py  │  │  personality_mapper.py      │            │
│  │  SQL Queries      │  │  MBTI / Holland / Numerology│            │
│  │  Caching          │  │  Cold-Start Mapping         │            │
│  └────────┬─────────┘  └────────────────────────────┘            │
│           │                                                       │
│  ┌────────▼─────────┐                                            │
│  │  PostgreSQL DB   │                                            │
│  │  (9study)        │                                            │
│  └──────────────────┘                                            │
└──────────────────────────────────────────────────────────────────┘
"""
add_code_block(doc, arch_text, "text")

doc.add_heading("Luồng xử lý chính (Main Flow)", level=2)
flow_steps = [
    "1. Client gửi POST request tới /recommend với thông tin học sinh (UserProfile)",
    "2. API Layer (api.py) validate input bằng Pydantic models",
    "3. DataLoader đã tải sẵn catalog (6,669 chương trình) khi server khởi động",
    "4. RecommendationEngine xử lý qua 3 giai đoạn: Filter → Score → Re-rank",
    "5. Kết quả trả về gồm: Top Majors Match, Recommendations, Safety Labels, Metadata",
]
for step in flow_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_page_break()


# ═══════════════════════════════════════════
# 3. CẤU TRÚC THƯ MỤC
# ═══════════════════════════════════════════

doc.add_heading("3. Cấu trúc thư mục & Mô tả file", level=1)

file_table = [
    ["api.py", "11,080 bytes / 319 dòng", "FastAPI REST API — Endpoint chính của hệ thống"],
    ["recommender.py", "20,285 bytes / 453 dòng", "Engine gợi ý — Pipeline 3 giai đoạn (Filter → Score → Re-rank)"],
    ["data_loader.py", "12,184 bytes / 283 dòng", "Tải và tiền xử lý dữ liệu từ PostgreSQL"],
    ["personality_mapper.py", "9,565 bytes / 175 dòng", "Ánh xạ MBTI/Holland/Numerology → Nhóm ngành & Từ khóa"],
    ["evaluate.py", "13,524 bytes / 343 dòng", "Đánh giá hệ thống bằng synthetic profiles"],
    ["demo.py", "5,173 bytes / 122 dòng", "Demo script với các test case mẫu"],
    ["db_introspection.py", "10,707 bytes / 261 dòng", "Khảo sát schema database PostgreSQL"],
    ["create_eda_notebook.py", "30,817 bytes / 703 dòng", "Tạo notebook EDA tự động"],
    ["EDA.ipynb", "~1.2 MB", "Jupyter Notebook phân tích dữ liệu khám phá"],
    ["requirements.txt", "161 bytes", "Danh sách thư viện phụ thuộc"],
    [".env", "192 bytes", "Biến môi trường (DATABASE_URL) — Không commit lên Git"],
    [".gitignore", "274 bytes", "Danh sách file/thư mục bị Git bỏ qua"],
    ["schema_snapshot.json", "~104 KB", "Snapshot schema database dạng JSON"],
]

add_styled_table(doc, ["Tên file", "Kích thước", "Mô tả"], file_table)

doc.add_page_break()


# ═══════════════════════════════════════════
# 4. CHI TIẾT CÁC MODULE
# ═══════════════════════════════════════════

doc.add_heading("4. Chi tiết các Module", level=1)

# 4.1 data_loader.py
doc.add_heading("4.1. data_loader.py — Module Tải & Tiền xử lý Dữ liệu", level=2)

doc.add_paragraph(
    "Module này chịu trách nhiệm kết nối tới PostgreSQL database, truy vấn dữ liệu từ các bảng "
    "chính (universities, programs, admission_scores, major_groups, program_admissions, subjects) "
    "và xây dựng Program Catalog — bộ dữ liệu tổng hợp dùng cho engine gợi ý."
)

doc.add_heading("Lớp DataLoader", level=3)
doc.add_paragraph("Các phương thức chính:")

dl_methods = [
    ["__init__(database_url)", "Khởi tạo kết nối DB, tạo SQLAlchemy engine, khởi tạo cache"],
    ["_read_sql_cached(query, key)", "Đọc SQL với in-memory caching — tránh truy vấn lặp"],
    ["load_universities()", "Tải danh sách trường ĐH (active) với metadata"],
    ["load_programs()", "Tải chương trình đào tạo, JOIN với universities"],
    ["load_major_groups()", "Tải nhóm ngành (id, name, icon)"],
    ["load_admission_scores()", "Tải điểm chuẩn, JOIN với programs và universities"],
    ["load_program_admissions()", "Tải phương thức xét tuyển cho từng chương trình"],
    ["load_subjects()", "Tải danh sách môn học"],
    ["build_program_catalog()", "XÂY DỰNG CATALOG — Phương thức chính, tạo DataFrame tổng hợp"],
    ["clear_cache()", "Xóa cache"],
]
add_styled_table(doc, ["Phương thức", "Mô tả"], dl_methods)

doc.add_heading("Quy trình build_program_catalog()", level=3)
catalog_steps = [
    "Tải programs + merge với major_groups để có tên nhóm ngành",
    "Lọc điểm THPT (method: thpt, hoc_ba, xthb, xt_hb), chỉ lấy score > 0 và ≤ 30, từ năm 2020+",
    "Tính latest_score_thpt: lấy median điểm chuẩn năm gần nhất cho mỗi chương trình",
    "Tính score_trend: so sánh 3 năm gần nhất → 'up' (tăng >1), 'down' (giảm >1), 'stable'",
    "Tổng hợp admission_methods và exam_blocks cho mỗi chương trình",
    "Parse exam_blocks từ JSONB sang Python list",
]
for i, step in enumerate(catalog_steps, 1):
    doc.add_paragraph(f"{i}. {step}", style='List Bullet')

doc.add_heading("SCORE_SCALE_MAP — Bản đồ thang điểm", level=3)
scale_data = [
    ["thpt", "0 – 30", "THPT Quốc gia (thang 30)"],
    ["hoc_ba", "0 – 30", "Xét học bạ (thang 30)"],
    ["dgnl", "0 – 1200", "Đánh giá năng lực"],
    ["dgtd", "0 – 100", "Đánh giá tư duy"],
    ["ccqt", "0 – 3000", "Chứng chỉ quốc tế"],
]
add_styled_table(doc, ["Phương thức", "Khoảng điểm", "Mô tả"], scale_data)

add_section_break(doc)

# 4.2 recommender.py
doc.add_heading("4.2. recommender.py — Engine Gợi ý (Core)", level=2)

doc.add_paragraph(
    "Đây là module cốt lõi của hệ thống, triển khai pipeline gợi ý 3 giai đoạn: "
    "Hard Filtering → Content-Based Scoring → Re-ranking & Safety Classification."
)

doc.add_heading("Lớp UserProfile (Dataclass)", level=3)
doc.add_paragraph("Đại diện cho hồ sơ đầu vào của học sinh, chia thành 3 cấp độ:")

profile_data = [
    ["Level 1 (Bắt buộc)", "exam_block, total_score, province", "Thông tin tối thiểu để gợi ý"],
    ["Level 2 (Khuyến nghị)", "subject_scores, budget_max, program_type, preferred_provinces, university_type", "Cải thiện độ chính xác"],
    ["Level 3 (Tuỳ chọn)", "interest_keywords, major_group_names, major_group_ids", "Cá nhân hóa sâu"],
    ["Personality", "holland_code, numerology, aspirations", "Cold-start & Blended Recommendation"],
]
add_styled_table(doc, ["Cấp độ", "Trường dữ liệu", "Mục đích"], profile_data)

doc.add_heading("Lớp RecommendationEngine", level=3)

doc.add_heading("Trọng số mặc định (DEFAULT_WEIGHTS)", level=4)
weight_data = [
    ["score_match", "0.30", "Mức độ khớp điểm thi với điểm chuẩn"],
    ["interest_match", "0.20", "Khớp sở thích/nhóm ngành"],
    ["location_match", "0.15", "Ưu tiên vị trí địa lý"],
    ["prestige_score", "0.10", "Uy tín trường (proxy)"],
    ["trend_bonus", "0.10", "Xu hướng điểm (stable/down = dễ hơn)"],
    ["budget_match", "0.10", "Phù hợp ngân sách"],
    ["quota_factor", "0.05", "Chỉ tiêu tuyển sinh (nhiều = dễ hơn)"],
]
add_styled_table(doc, ["Yếu tố", "Trọng số", "Ý nghĩa"], weight_data)

doc.add_heading("Stage 1: Hard Filtering (_hard_filter)", level=4)
doc.add_paragraph(
    "Loại bỏ các chương trình không khả thi dựa trên 2 tiêu chí cứng:"
)
filter_items = [
    "Exam Block Compatibility: Chỉ giữ chương trình chấp nhận khối thi của user. "
    "Nếu chương trình không có thông tin khối thi, mặc định giữ lại.",
    "Score Feasibility: Loại chương trình có điểm chuẩn > (điểm user + margin). "
    "Margin mặc định = 3.0 điểm, cho phép gợi ý 'Reach' (thử thách).",
]
for item in filter_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("Stage 2: Content-Based Scoring (_compute_scores)", level=4)
doc.add_paragraph("Hệ thống tính 7 sub-scores cho mỗi chương trình, sau đó kết hợp bằng trọng số:")

scoring_funcs = [
    ["_score_match()", "Gaussian similarity: exp(-0.5*(diff/σ)²), σ=3.0. Nếu user > cutoff: bonus safety. Nếu không có điểm: trả 0.4"],
    ["_location_match()", "1.0 nếu cùng tỉnh, 0.75 nếu trong preferred, 0.4 nếu thành phố lớn, 0.2 nếu khác"],
    ["_budget_match()", "1.0 nếu trong ngân sách, giảm dần nếu vượt. Không có info: 0.7"],
    ["_interest_match()", "So khớp major_group_names và keyword search trong search_text"],
    ["_quota_factor()", "Tỷ lệ quota/max_quota, chuẩn hóa về [0.3, 0.8]"],
    ["_prestige_score()", "Proxy từ is_featured, official_link, điểm chuẩn > 20"],
    ["_trend_bonus()", "down=0.7, stable=0.6, up=0.4, unknown=0.5"],
]
add_styled_table(doc, ["Hàm", "Công thức / Logic"], scoring_funcs)

doc.add_heading("Stage 3: Re-ranking & Safety Classification (_rerank)", level=4)
doc.add_paragraph("Phân loại mức độ an toàn cho mỗi đề xuất:")

safety_data = [
    ["safe (An toàn)", "User score ≥ Cutoff + 1", "Khả năng đậu cao"],
    ["match (Vừa sức)", "User score ± 1 so với Cutoff", "Cơ hội tốt, cần cố gắng"],
    ["reach (Thử thách)", "User score từ -3 đến -1 so với Cutoff", "Khó nhưng có thể"],
    ["ambitious (Tham vọng)", "User score < Cutoff - 3", "Rất khó, cần nỗ lực đặc biệt"],
]
add_styled_table(doc, ["Mức độ", "Điều kiện", "Ý nghĩa"], safety_data)

doc.add_paragraph(
    "Re-ranking áp dụng ràng buộc đa dạng: tối đa 3 chương trình/trường, "
    "tối đa 5 chương trình/nhóm ngành. Nguyện vọng (aspirations) được ưu tiên "
    "và bỏ qua các ràng buộc đa dạng."
)

add_section_break(doc)

# 4.3 personality_mapper.py
doc.add_heading("4.3. personality_mapper.py — Ánh xạ Tính cách", level=2)

doc.add_paragraph(
    "Module giải quyết vấn đề Cold-Start (khi user chưa biết thích ngành gì) "
    "bằng cách ánh xạ từ trắc nghiệm tính cách sang nhóm ngành và từ khóa sở thích."
)

doc.add_heading("Hệ thống ánh xạ", level=3)

mapper_data = [
    ["Holland RIASEC", "6 mã (R/I/A/S/E/C)", "Mỗi mã → 4-6 nhóm ngành + 6 keywords"],
    ["MBTI", "16 loại tính cách", "Mỗi loại → 3 mã Holland + bonus groups"],
    ["Thần số học", "12 số (1-9, 11, 22, 33)", "Mỗi số → 4 nhóm ngành phù hợp"],
]
add_styled_table(doc, ["Hệ thống", "Đầu vào", "Đầu ra"], mapper_data)

doc.add_heading("Holland RIASEC Map (chi tiết)", level=3)
holland_data = [
    ["R - Realistic", "Kỹ thuật ô tô, Điện - Điện tử, Xây dựng, Nông - Lâm, Vật liệu", "kỹ thuật, cơ khí, xây dựng"],
    ["I - Investigative", "KHMT, Y - Dược, Khoa học tự nhiên, CNSH, Toán - Thống kê", "nghiên cứu, khoa học, dữ liệu"],
    ["A - Artistic", "Nghệ thuật, Ngôn ngữ, Văn học, Dệt may, Kiến trúc", "thiết kế, sáng tạo, nghệ thuật"],
    ["S - Social", "Sư phạm, Tâm lý, Y - Dược, Du lịch, Quản lý nhân lực", "giáo dục, tâm lý, xã hội"],
    ["E - Enterprising", "Kinh tế - QTKD, Marketing, Luật, KD quốc tế, Tài chính", "kinh doanh, quản trị, marketing"],
    ["C - Conventional", "Kế toán, Tài chính, Quản lý nhà nước, Toán - Thống kê", "kế toán, tài chính, số liệu"],
]
add_styled_table(doc, ["Mã Holland", "Nhóm ngành chính", "Từ khóa"], holland_data)

doc.add_heading("Top Major Match — Công thức tính phần trăm khớp", level=3)
doc.add_paragraph(
    "Phần trăm khớp cho mỗi nhóm ngành được tính dựa trên 3 nguồn:\n"
    "• Holland Match (trọng số 50%): Primary letter = 1.0, Secondary = 0.8, Tertiary = 0.6\n"
    "• Numerology Match (30%): Binary (có/không có trong danh sách)\n"
    "• Explicit Interest (20%): User đã chọn nhóm ngành hoặc aspirations liên quan\n"
    "\nCông thức kết hợp: final_score = holland × 0.5 + numerology × 0.3 + interest × 0.2"
)

add_section_break(doc)

# 4.4 api.py
doc.add_heading("4.4. api.py — FastAPI REST API", level=2)

doc.add_paragraph(
    "Module cung cấp REST API cho hệ thống gợi ý. Sử dụng FastAPI với Pydantic validation, "
    "CORS middleware, và Swagger UI tự động."
)

doc.add_heading("Pydantic Models (Request/Response)", level=3)
model_data = [
    ["RecommendationRequest", "Input chính — 14 trường (3 bắt buộc, 11 tùy chọn)"],
    ["RecommendationResponse", "Output — top_majors, recommendations, metadata, explanations"],
    ["ProgramRecommendation", "Chi tiết 1 chương trình được gợi ý"],
    ["PersonalityRequest", "Input cho personality mapping (MBTI/Holland)"],
    ["PersonalityResponse", "Output — major_group_names, interest_keywords"],
    ["HealthResponse", "Trạng thái hệ thống — status, catalog_size, programs_with_scores"],
    ["MajorGroupResponse", "Thông tin 1 nhóm ngành (id, name, icon)"],
]
add_styled_table(doc, ["Model", "Mô tả"], model_data)

doc.add_heading("Lifecycle Management", level=3)
doc.add_paragraph(
    "Server sử dụng asynccontextmanager (lifespan) để:\n"
    "• Startup: Tải DataLoader, xây dựng catalog, khởi tạo RecommendationEngine\n"
    "• Shutdown: Giải phóng tài nguyên (clear engine_state)\n"
    "• Global state lưu trong dict engine_state"
)

add_section_break(doc)

# 4.5 evaluate.py
doc.add_heading("4.5. evaluate.py — Đánh giá Hệ thống", level=2)

doc.add_paragraph(
    "Module chạy đánh giá tự động với 8 synthetic test profiles (hồ sơ giả lập) "
    "đại diện cho các kịch bản sử dụng khác nhau."
)

doc.add_heading("Danh sách Synthetic Profiles", level=3)
profile_table = [
    ["Khá Tự Nhiên - Hà Nội", "A00/A01, 24.0 điểm", "KHMT - Kỹ thuật PM"],
    ["Giỏi Xã Hội - TP.HCM", "D01/C00, 27.0 điểm", "Kinh tế - QTKD"],
    ["Trung Bình - Đà Nẵng", "A00, 18.5 điểm, Budget 25tr", "Bất kỳ (budget test)"],
    ["Xuất sắc - Y Dược", "B00, 28.5 điểm", "Y - Dược"],
    ["Khá - Sư Phạm - Tỉnh lẻ", "A00/A01/D01, 22.0 điểm", "Sư phạm"],
    ["Cold Start - Level 1 only", "A00, 20.0 điểm", "Bất kỳ"],
    ["Giỏi - Ngôn Ngữ Anh", "D01, 26.0 điểm", "Ngôn ngữ - Ngoại ngữ"],
    ["Trung Bình Khá - Xây Dựng", "A00/A01, 21.0 điểm", "Xây dựng - Kiến trúc"],
]
add_styled_table(doc, ["Tên profile", "Thông tin thi", "Ngành mong đợi"], profile_table)

doc.add_heading("Metrics đánh giá", level=3)
metrics_data = [
    ["Interest Precision", "% kết quả thuộc nhóm ngành mong đợi", "≥ 40%"],
    ["Location Precision", "% kết quả đúng tỉnh/TP mong đợi", "≥ 50%"],
    ["Score Feasibility", "% kết quả user có thể đậu (score + margin)", "≥ 90%"],
    ["Program Coverage", "% chương trình được gợi ý ít nhất 1 lần", "≥ 5%"],
    ["University Coverage", "% trường được gợi ý ít nhất 1 lần", "≥ 10%"],
    ["Avg Unique Universities", "Trung bình số trường khác nhau/profile", "≥ 3"],
    ["Avg Unique Major Groups", "Trung bình số nhóm ngành khác nhau/profile", "≥ 2"],
]
add_styled_table(doc, ["Metric", "Ý nghĩa", "Target"], metrics_data)

add_section_break(doc)

# 4.6 demo.py
doc.add_heading("4.6. demo.py — Script Demo", level=2)
doc.add_paragraph(
    "Script chạy nhanh 4 test case để kiểm tra engine gợi ý. Các test case bao gồm:\n"
    "1. Học sinh khá, khối A00, Hà Nội — hướng CNTT\n"
    "2. Học sinh giỏi, khối D01, TP.HCM, budget hạn chế — hướng Kinh tế\n"
    "3. Cold Start — chỉ có thông tin Level 1\n"
    "4. Học sinh điểm rất cao — hướng Y Dược"
)

add_section_break(doc)

# 4.7 db_introspection.py
doc.add_heading("4.7. db_introspection.py — Khảo sát Database", level=2)
doc.add_paragraph(
    "Script Phase 1 — kết nối PostgreSQL, sử dụng SQLAlchemy Inspector để tự động:\n"
    "• Quét toàn bộ schema (tables, columns, types, PK, FK)\n"
    "• Xây dựng Entity-Relationship Map\n"
    "• Đánh giá Feature Readiness cho hệ thống gợi ý\n"
    "• Export schema ra file schema_snapshot.json\n\n"
    "Tất cả operations đều READ-ONLY, không thay đổi dữ liệu."
)

add_section_break(doc)

# 4.8 create_eda_notebook.py
doc.add_heading("4.8. create_eda_notebook.py — Tạo Notebook EDA", level=2)
doc.add_paragraph(
    "Script Phase 2 — tự động sinh file EDA.ipynb (Jupyter Notebook) với 23 cells bao gồm:\n"
    "• Setup & Database Connection\n"
    "• Load 7 bảng dữ liệu chính\n"
    "• Data Quality Profiling (Missing Values Heatmap)\n"
    "• Distribution Analysis (Scores, Provinces, Major Groups)\n"
    "• Outlier Detection\n"
    "• Class Imbalance Analysis\n"
    "• Feature Correlation (Score vs Tuition, Score vs Quota)\n"
    "• Score Trends Over Years\n"
    "• Summary Report"
)

doc.add_page_break()


# ═══════════════════════════════════════════
# 5. MÔ HÌNH DỮ LIỆU
# ═══════════════════════════════════════════

doc.add_heading("5. Mô hình dữ liệu (Data Model)", level=1)

doc.add_paragraph(
    "Hệ thống sử dụng PostgreSQL database '9study' với các bảng chính sau:"
)

er_text = """
universities ─────┐
  ├── id (PK)     │
  ├── university_code  
  ├── name        │     programs ────────────┐
  ├── province    │       ├── id (PK)        │
  ├── type        ├──FK── ├── university_id  │
  ├── website     │       ├── major_code     │
  ├── tuition_*   │       ├── major_name     │
  └── ...         │       ├── major_group_id ──FK── major_groups
                  │       ├── program_type   │       ├── id (PK)
                  │       ├── tuition_min    │       ├── name
                  │       ├── tuition_max    │       └── icon
                  │       └── total_quota    │
                  │                          │
                  │     program_admissions ──┤
                  │       ├── id (PK)        │
                  │       ├── program_id ──FK┘
                  │       ├── method
                  │       └── exam_blocks (JSONB)
                  │              │
                  │     admission_scores
                  │       ├── id (PK)
                  │       ├── program_admission_id ──FK
                  │       ├── year
                  │       ├── score
                  │       └── note
                  │
                  │     subjects
                  │       ├── id (PK)
                  │       ├── name
                  │       └── short_name
"""
add_code_block(doc, er_text, "text")

doc.add_heading("Program Catalog (Output của DataLoader)", level=2)
catalog_cols = [
    ["program_id", "UUID", "ID chương trình"],
    ["university_code", "String", "Mã trường (e.g. BKA)"],
    ["university_name", "String", "Tên trường đầy đủ"],
    ["province", "String", "Tỉnh/TP của trường"],
    ["university_type", "String", "Công lập / Dân lập"],
    ["major_code", "String", "Mã ngành"],
    ["major_name", "String", "Tên ngành"],
    ["major_group_name", "String", "Tên nhóm ngành"],
    ["program_type", "String", "Đại trà / Chất lượng cao / ..."],
    ["tuition_min / tuition_max", "Float", "Học phí min/max (VND/năm)"],
    ["total_quota", "Float", "Chỉ tiêu tuyển sinh"],
    ["latest_score_thpt", "Float", "Điểm chuẩn THPT năm gần nhất"],
    ["latest_score_year", "Float", "Năm của điểm chuẩn"],
    ["score_trend", "String", "Xu hướng: up / down / stable / unknown"],
    ["admission_methods", "List", "Danh sách phương thức xét tuyển"],
    ["exam_blocks_list", "List", "Danh sách khối thi chấp nhận"],
]
add_styled_table(doc, ["Cột", "Kiểu", "Mô tả"], catalog_cols)

doc.add_page_break()


# ═══════════════════════════════════════════
# 6. THUẬT TOÁN GỢI Ý
# ═══════════════════════════════════════════

doc.add_heading("6. Thuật toán Gợi ý — Pipeline 3 Giai đoạn", level=1)

doc.add_heading("Tổng quan Pipeline", level=2)
pipeline_text = """
Input: UserProfile (exam_block, total_score, province, ...)
                │
        ┌───────▼──────────┐
        │  Top Majors Match │  ← Tính % khớp cho từng nhóm ngành
        │  (Holland + Num)  │     trước khi filter
        └───────┬──────────┘
                │
     ┌──────────▼──────────────┐
     │  STAGE 1: Hard Filter    │  catalog: 6,669 → ~1,000-3,000
     │  • Exam block match      │
     │  • Score feasibility     │
     └──────────┬──────────────┘
                │
     ┌──────────▼──────────────┐
     │  STAGE 2: Scoring        │  7 sub-scores × weights
     │  • Score match (0.30)    │
     │  • Interest match (0.20) │
     │  • Location match (0.15) │
     │  • Prestige (0.10)       │
     │  • Trend bonus (0.10)    │
     │  • Budget match (0.10)   │
     │  • Quota factor (0.05)   │
     └──────────┬──────────────┘
                │
     ┌──────────▼──────────────┐
     │  STAGE 3: Re-rank        │  Safety labels + Diversity
     │  • Safety classification │
     │  • Max 3/university      │
     │  • Max 5/major group     │
     │  • Aspirations bypass    │
     └──────────┬──────────────┘
                │
Output: { top_majors, recommendations, metadata }
"""
add_code_block(doc, pipeline_text, "text")

doc.add_heading("Score Match — Chi tiết hàm Gaussian Similarity", level=2)
doc.add_paragraph(
    "Hàm _score_match() sử dụng Gaussian similarity để đo mức độ khớp giữa "
    "điểm user và điểm chuẩn chương trình:"
)
score_formula = """
base = exp(-0.5 × ((user_score - program_score) / σ)²)    [σ = 3.0]

Nếu user_score > program_score (user cao hơn cutoff):
    safety_bonus = min(0.15, diff × 0.03)
    final = min(1.0, base + safety_bonus)

Nếu user_score ≤ program_score (user thấp hơn cutoff):
    final = base × 0.9    [phạt nhẹ]

Nếu không có điểm chuẩn:
    final = 0.4            [default trung bình]
"""
add_code_block(doc, score_formula, "text")

doc.add_heading("Blended Recommendation — Kết hợp Aspiration & Discovery", level=2)
doc.add_paragraph(
    "Hệ thống kết hợp 2 loại kết quả:\n"
    "1. Aspiration (Nguyện vọng): Chương trình user chủ động đăng ký — được ưu tiên, bypass diversity constraints\n"
    "2. Discovery (Khám phá): Chương trình hệ thống tự tìm — tuân theo diversity constraints\n\n"
    "Khi user có Holland/Numerology, hệ thống tự mở rộng interest_keywords "
    "bằng các nhóm ngành có match ≥ 60% trước khi filter, giúp khám phá thêm ngành phù hợp."
)

doc.add_page_break()


# ═══════════════════════════════════════════
# 7. API REFERENCE
# ═══════════════════════════════════════════

doc.add_heading("7. API Reference", level=1)

doc.add_paragraph(
    "Server chạy tại http://localhost:8000 | Swagger UI tại /docs"
)

api_endpoints = [
    ["GET", "/", "Redirect tới /docs (Swagger UI)", "—"],
    ["GET", "/health", "Health check — catalog size, score count", "HealthResponse"],
    ["GET", "/major-groups", "Danh sách tất cả nhóm ngành", "List[MajorGroupResponse]"],
    ["POST", "/onboarding/personality", "Ánh xạ MBTI/Holland → nhóm ngành", "PersonalityResponse"],
    ["POST", "/recommend", "GỢI Ý CHÍNH — Nhận profile, trả recommendations", "RecommendationResponse"],
]
add_styled_table(doc, ["Method", "Path", "Mô tả", "Response Model"], api_endpoints)

doc.add_heading("Ví dụ Request — POST /recommend", level=2)
example_req = """{
    "exam_block": ["A00"],
    "total_score": 24.5,
    "province": "Hà Nội",
    "preferred_provinces": ["Hà Nội", "TP.HCM"],
    "interest_keywords": ["công nghệ", "phần mềm"],
    "major_group_names": ["Khoa học máy tính - Kỹ thuật phần mềm"],
    "top_k": 10,
    "score_margin": 3.0
}"""
add_code_block(doc, example_req, "json")

doc.add_heading("Ví dụ Response (rút gọn)", level=2)
example_resp = """{
    "top_majors": [
        {"major_group_name": "Khoa học máy tính - Kỹ thuật PM", "match_percentage": 85}
    ],
    "recommendations": [
        {
            "program_id": "uuid-...",
            "university_code": "BKA",
            "university_name": "Đại học Bách khoa Hà Nội",
            "major_name": "Khoa học máy tính",
            "province": "Hà Nội",
            "latest_score_thpt": 25.5,
            "safety_level": "match",
            "recommendation_score": 0.7823,
            "is_aspiration": false
        }
    ],
    "metadata": {
        "total_programs": 6669,
        "after_filter": 2100,
        "returned": 10,
        "safety_distribution": {"safe": 4, "match": 3, "reach": 2, "ambitious": 1}
    }
}"""
add_code_block(doc, example_resp, "json")

doc.add_page_break()


# ═══════════════════════════════════════════
# 8. CẤU HÌNH & PHỤ THUỘC
# ═══════════════════════════════════════════

doc.add_heading("8. Cấu hình & Phụ thuộc (Dependencies)", level=1)

doc.add_heading("Biến môi trường (.env)", level=2)
env_data = [
    ["DATABASE_URL", "postgresql://user:pass@host:port/dbname", "Bắt buộc", "Kết nối PostgreSQL"],
]
add_styled_table(doc, ["Biến", "Định dạng", "Yêu cầu", "Mục đích"], env_data)

doc.add_heading("Thư viện Python (requirements.txt)", level=2)
deps_data = [
    ["pandas", "≥ 2.0.0", "Xử lý dữ liệu dạng bảng (DataFrame)"],
    ["psycopg2-binary", "≥ 2.9.0", "Driver PostgreSQL cho Python"],
    ["sqlalchemy", "≥ 2.0.0", "ORM và SQL toolkit"],
    ["python-dotenv", "≥ 1.0.0", "Đọc biến môi trường từ .env"],
    ["jupyter", "≥ 1.0.0", "Chạy EDA notebook"],
    ["seaborn", "≥ 0.13.0", "Trực quan hóa thống kê"],
    ["scikit-learn", "≥ 1.3.0", "Machine learning utilities"],
    ["matplotlib", "≥ 3.7.0", "Trực quan hóa biểu đồ"],
    ["tabulate", "≥ 0.9.0", "Format bảng trong terminal"],
    ["fastapi", "latest", "Web framework (khai báo riêng)"],
    ["uvicorn", "latest", "ASGI server cho FastAPI"],
]
add_styled_table(doc, ["Thư viện", "Phiên bản", "Mục đích"], deps_data)

doc.add_heading("Cách khởi chạy", level=2)
run_commands = """
# 1. Cài đặt dependencies
pip install -r requirements.txt
pip install fastapi uvicorn

# 2. Cấu hình database
# Tạo file .env với DATABASE_URL

# 3. Chạy API server
uvicorn api:app --reload --port 8000

# 4. Mở Swagger UI
# Truy cập: http://localhost:8000/docs

# 5. Chạy demo
python demo.py

# 6. Chạy evaluation
python evaluate.py
"""
add_code_block(doc, run_commands, "bash")

doc.add_page_break()


# ═══════════════════════════════════════════
# 9. ĐÁNH GIÁ & METRICS
# ═══════════════════════════════════════════

doc.add_heading("9. Đánh giá & Metrics", level=1)

doc.add_paragraph(
    "Hệ thống được đánh giá tự động bằng module evaluate.py với 8 synthetic test profiles, "
    "kiểm tra trên 5 nhóm metrics chính."
)

doc.add_heading("Thống kê dữ liệu (tại thời điểm chạy)", level=2)
stats_data = [
    ["Tổng chương trình (catalog)", "6,669"],
    ["Chương trình có điểm THPT", "3,615 (54.2%)"],
    ["Tổng nhóm ngành", "35"],
    ["Tổng trường đại học", "~200+"],
]
add_styled_table(doc, ["Metric", "Giá trị"], stats_data)

doc.add_heading("Tiêu chí đánh giá PASS/FAIL", level=2)
criteria_data = [
    ["Average Interest Precision", "≥ 40%", "Kết quả đúng nhóm ngành mong đợi"],
    ["Average Score Feasibility", "≥ 90%", "Kết quả user có thể đậu"],
    ["Overall", "Interest ≥ 30% AND Feasibility ≥ 90%", "Hệ thống hoạt động đúng"],
]
add_styled_table(doc, ["Tiêu chí", "Ngưỡng", "Ý nghĩa"], criteria_data)

doc.add_page_break()


# ═══════════════════════════════════════════
# 10. MÃ NGUỒN ĐẦY ĐỦ
# ═══════════════════════════════════════════

doc.add_heading("10. Mã nguồn đầy đủ (Full Source Code)", level=1)

doc.add_paragraph(
    "Phần này chứa toàn bộ mã nguồn của các file chính trong dự án, "
    "được trình bày nguyên văn để tham khảo."
)

# Order of files to include
source_order = [
    ("api.py", "FastAPI REST API — Endpoint chính"),
    ("recommender.py", "Recommendation Engine — Core Algorithm"),
    ("data_loader.py", "Data Loader — Tải & Tiền xử lý Dữ liệu"),
    ("personality_mapper.py", "Personality Mapper — MBTI/Holland/Numerology"),
    ("evaluate.py", "Evaluation Suite — Đánh giá Hệ thống"),
    ("demo.py", "Demo Script — Test Cases"),
    ("db_introspection.py", "Database Introspection — Khảo sát Schema"),
    ("requirements.txt", "Dependencies"),
]

for filename, description in source_order:
    doc.add_heading(f"10.x. {filename}", level=2)
    p = doc.add_paragraph()
    run = p.add_run(f"📄 {description}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x5A, 0x6A, 0x7A)
    
    code = source_files.get(filename, "[Not found]")
    
    # Split into chunks if too long (Word has limits)
    lines = code.split('\n')
    chunk_size = 80
    for start in range(0, len(lines), chunk_size):
        chunk = '\n'.join(lines[start:start + chunk_size])
        add_code_block(doc, chunk)
    
    if filename != source_order[-1][0]:
        doc.add_page_break()


# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════

output_path = os.path.join(PROJECT_DIR, "9study_Source_Code_Report.docx")
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
print(f"   Total pages: ~50+ (estimated)")
